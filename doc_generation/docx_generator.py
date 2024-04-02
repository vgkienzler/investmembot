import os
from pprint import pprint
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(document_structure: dict):
    """ Generate a .docx document from the json document structure.

    :arg
    document_structure: dict: The structure of the document to generate.
    """
    # Create a new Document
    doc = Document()

    print('Document structure: ')
    pprint(document_structure)

    # Add a title
    title = doc.add_heading(document_structure['document_title'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section in document_structure['sections']:
        # Extract title and content from the document
        if section['type'] == 'title':
            title = section['content']
            title_level = section['level']

            # Add to the Word document
            title = doc.add_heading(title, level=title_level)
            title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif section['type'] == 'text':
            content = section['content']
            print("Content as seen by generate_docx(): ", content)

            # Add to the Word document
            p = doc.add_paragraph()
            p.add_run(content)

    try:
        output = os.getenv('OUTPUT_FOLDER_PATH', 'output')
        doc.save(Path(output) / 'demo.docx')
    except PermissionError:
        print('A document with the same name already exists, pleased delete it before running '
              'this script again.')

    return doc


if __name__ == '__main__':
    from doc_generator_test_structure import doc_structure_test

    generate_docx(doc_structure_test)
